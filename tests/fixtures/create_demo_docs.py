#!/usr/bin/env python3
"""
Create demo Word document for testing MCP tool
"""

try:
    from docx import Document

    # Create Word document
    doc = Document()
    doc.add_heading('Model Context Protocol (MCP)', 0)

    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'The Model Context Protocol (MCP) is an open protocol that enables '
        'seamless integration between LLM applications and external data sources '
        'and tools.'
    )

    doc.add_heading('Key Features', level=1)
    doc.add_paragraph('MCP provides three main primitives:')

    # Add bullet points
    doc.add_paragraph('Resources - Access to data and content', style='List Bullet')
    doc.add_paragraph('Tools - Functions that the LLM can invoke', style='List Bullet')
    doc.add_paragraph('Prompts - Pre-defined prompt templates', style='List Bullet')

    doc.add_heading('Benefits', level=1)
    doc.add_paragraph(
        'MCP enables developers to build more capable AI applications by '
        'giving models access to the right context at the right time. '
        'This includes data from various sources, tools to interact with systems, '
        'and standardized ways to structure prompts.'
    )

    # Save document
    doc.save('/home/user/training/tests/fixtures/mcp_demo.docx')
    print("✅ Created mcp_demo.docx")

except ImportError:
    print("❌ python-docx not installed. Run: pip install python-docx")
except Exception as e:
    print(f"❌ Error creating document: {e}")
