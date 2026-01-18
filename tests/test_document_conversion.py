#!/usr/bin/env python3
"""
Test suite for document_path_to_markdown tool
Following TDD approach from Claude Code lesson
"""

import pytest
import os
import sys
import tempfile
from pathlib import Path

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

# Import the document conversion functionality
from document_utils import document_path_to_markdown


class TestDocumentPathToMarkdown:
    """Test suite for document path to markdown conversion"""

    def setup_method(self):
        """Set up test fixtures"""
        self.test_data_dir = Path(__file__).parent.parent / "data"
        self.temp_dir = tempfile.mkdtemp()

    def test_valid_pdf_conversion(self):
        """Test 1: Verify PDF files are correctly converted to Markdown"""
        # Look for existing PDF in data directory
        pdf_files = list(self.test_data_dir.glob("*.pdf"))

        if pdf_files:
            pdf_path = str(pdf_files[0])
            result = document_path_to_markdown(pdf_path)

            # Verify result is a string
            assert isinstance(result, str)
            # Verify result is not empty
            assert len(result) > 0
            # Verify it contains some text content
            assert result.strip() != ""
        else:
            pytest.skip("No PDF files available for testing")

    def test_valid_docx_conversion(self):
        """Test 2: Verify Word documents are correctly converted to Markdown"""
        # Create a simple test DOCX file
        docx_path = os.path.join(self.temp_dir, "test.docx")

        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Test Document', 0)
            doc.add_paragraph('This is a test paragraph.')
            doc.save(docx_path)

            result = document_path_to_markdown(docx_path)

            # Verify result is a string
            assert isinstance(result, str)
            # Verify it contains the test content
            assert "Test Document" in result or "test" in result.lower()
            assert len(result) > 0
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_non_existent_file_error(self):
        """Test 3: Verify appropriate error for missing files"""
        non_existent_path = "/path/to/nonexistent/file.pdf"

        with pytest.raises(FileNotFoundError) as exc_info:
            document_path_to_markdown(non_existent_path)

        # Verify error message is descriptive
        assert "not found" in str(exc_info.value).lower() or "does not exist" in str(exc_info.value).lower()

    def test_unsupported_file_type_error(self):
        """Test 4: Verify error for unsupported file formats"""
        # Create a test .txt file
        txt_path = os.path.join(self.temp_dir, "test.txt")
        with open(txt_path, 'w') as f:
            f.write("This is a text file")

        with pytest.raises(ValueError) as exc_info:
            document_path_to_markdown(txt_path)

        # Verify error message mentions unsupported format
        assert "unsupported" in str(exc_info.value).lower() or "format" in str(exc_info.value).lower()

    def test_markdown_output_format(self):
        """Test 5: Verify the output is valid Markdown format"""
        # Use an existing PDF if available
        pdf_files = list(self.test_data_dir.glob("*.pdf"))

        if pdf_files:
            pdf_path = str(pdf_files[0])
            result = document_path_to_markdown(pdf_path)

            # Basic Markdown validation
            # Should be plain text (no binary data)
            assert isinstance(result, str)
            # Should not contain null bytes
            assert '\x00' not in result
            # Should be decodable as text
            try:
                result.encode('utf-8')
                assert True
            except UnicodeEncodeError:
                assert False, "Result contains non-text characters"
        else:
            pytest.skip("No PDF files available for testing")


if __name__ == "__main__":
    # Run tests
    pytest.main([__file__, "-v"])
